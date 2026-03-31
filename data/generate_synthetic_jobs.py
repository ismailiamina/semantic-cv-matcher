"""
Générateur de Fiches de Poste Synthétiques — Format DOCX

"""

import json
import random
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


ENTREPRISES = [
    "OCP", "Maroc Telecom", "BMCE Bank", "Attijariwafa Bank",
    "CIH Bank", "TotalEnergies Maroc", "Renault Maroc",
    "Société Générale Maroc", "BNP Paribas Maroc", "Orange Maroc",
    "Inwi", "HPS", "CGI Maroc", "Capgemini Maroc", "Accenture Maroc",
    "Deloitte Maroc", "Ernst & Young Maroc", "Leyton Maroc",
    "Alten Maroc", "Sopra Steria Maroc", "IBM Maroc", "Oracle Maroc"
]

VILLES = [
    "Casablanca", "Rabat", "Marrakech", "Fès", "Tanger",
    "Agadir", "Meknès", "Kenitra", "Remote"
]

SECTEURS = [
    "Finance", "Banque", "Assurance", "E-commerce", "Santé",
    "Télécommunications", "Énergie", "Logistique", "Éducation",
    "Immobilier", "Transport", "Retail", "Manufacturing", "IT Services"
]

TYPES_EMPLOI = ["Full-time", "Part-time", "Fixed-term", "Internship"]

NIVEAUX_EXPERIENCE = ["Junior", "Medior", "Confirmé", "Senior", "Expert"]

NIVEAUX_EDUCATION = ["Bac + 2", "Bac + 3", "Bac + 4", "Bac + 5", "Bac + 6 et plus"]

CERTIFICATIONS = [
    "AWS Solutions Architect", "AWS Cloud Practitioner",
    "Google Cloud Professional", "Azure Administrator",
    "Kubernetes CKA", "Terraform Associate",
    "Scrum Master PSM I", "PMP", "CISSP"
]

# Jobs par domaine
JOBS = {
    "data": {
        "titles": [
            "Data Engineer", "Data Scientist", "ML Engineer",
            "MLOps Engineer", "Data Analyst", "Lead Data Engineer",
            "Senior Data Scientist", "AI Engineer"
        ],
        "tech_skills": [
            "Apache Spark", "Airflow", "Kafka", "MLflow", "Kubernetes",
            "Docker", "TensorFlow", "PyTorch", "Scikit-learn", "Pandas",
            "Power BI", "Tableau", "Elasticsearch", "Redis", "FastAPI",
            "AWS SageMaker", "Databricks", "dbt", "Great Expectations"
        ],
        "prog_langs": ["Python", "SQL", "Scala", "R", "Julia"],
        "industries": ["Finance", "Banque", "Santé", "E-commerce", "Télécommunications"]
    },
    "web": {
        "titles": [
            "Développeur Full Stack", "Développeur Backend", "Développeur Frontend",
            "Ingénieur API", "Tech Lead", "Architecte Logiciel",
            "Senior Developer Node.js", "Lead Frontend Engineer"
        ],
        "tech_skills": [
            "React", "Angular", "Vue.js", "Node.js", "Django", "FastAPI",
            "Spring Boot", "Docker", "Kubernetes", "REST API", "GraphQL",
            "PostgreSQL", "MongoDB", "Redis", "Nginx", "Jenkins",
            "TypeScript", "Tailwind CSS", "Microservices"
        ],
        "prog_langs": ["JavaScript", "TypeScript", "Python", "Java", "PHP", "Go"],
        "industries": ["E-commerce", "Fintech", "Éducation", "Retail", "IT Services"]
    },
    "devops": {
        "titles": [
            "DevOps Engineer", "SRE Engineer", "Cloud Architect",
            "Infrastructure Engineer", "Platform Engineer",
            "Senior DevOps Engineer", "Lead Cloud Engineer"
        ],
        "tech_skills": [
            "Docker", "Kubernetes", "Terraform", "Ansible", "Jenkins",
            "GitLab CI/CD", "AWS", "Azure", "GCP", "Prometheus",
            "Grafana", "ELK Stack", "Helm", "ArgoCD", "Linux", "Bash"
        ],
        "prog_langs": ["Python", "Bash", "Go", "YAML", "HCL"],
        "industries": ["Banque", "Télécommunications", "IT Services", "Finance"]
    },
    "security": {
        "titles": [
            "Ingénieur Cybersécurité", "Pentester", "Security Analyst",
            "SOC Analyst", "RSSI", "Senior Security Engineer"
        ],
        "tech_skills": [
            "SIEM", "Splunk", "Wireshark", "Metasploit", "Burp Suite",
            "OWASP", "ISO 27001", "Firewall", "IDS/IPS", "PKI",
            "Kali Linux", "Nessus", "CrowdStrike", "Azure Sentinel"
        ],
        "prog_langs": ["Python", "Bash", "C", "PowerShell"],
        "industries": ["Banque", "Assurance", "Énergie", "Télécommunications"]
    }
}



def get_experience_years(level: str) -> int:
    """Retourne les années d'expérience selon le niveau"""
    mapping = {
        "Junior": random.randint(0, 2),
        "Medior": random.randint(2, 4),
        "Confirmé": random.randint(4, 6),
        "Senior": random.randint(6, 10),
        "Expert": random.randint(10, 15)
    }
    return mapping.get(level, 3)


def generate_job_description(title, tech_skills, prog_langs, level, company, industry):
    """Génère une description de poste réaliste"""
    skills_str = ", ".join(tech_skills[:5])
    langs_str = ", ".join(prog_langs[:3])

    descriptions = [
        f"{company} recherche un(e) {title} {level} pour rejoindre son équipe technique "
        f"dans le secteur {industry}. "
        f"Vous serez responsable de la conception et du développement de solutions innovantes. "
        f"Compétences requises : {skills_str}. "
        f"Langages : {langs_str}. "
        f"Vous travaillerez en méthodologie Agile/Scrum avec une équipe internationale.",

        f"Dans le cadre de son développement, {company} recrute un(e) {title} "
        f"pour renforcer son pôle {industry}. "
        f"Missions principales : développement et maintenance des systèmes, "
        f"participation aux revues de code, optimisation des performances. "
        f"Technologies utilisées : {skills_str}. Langages : {langs_str}.",

        f"Rejoignez {company}, acteur majeur du secteur {industry}, "
        f"en tant que {title} {level}. "
        f"Vous contribuerez à des projets à fort impact technologique. "
        f"Stack technique : {skills_str}. "
        f"Programmation : {langs_str}. "
        f"Environnement de travail stimulant et opportunités d'évolution."
    ]

    return random.choice(descriptions)


def generate_summary(title, company, level, skills, years):
    """Génère un résumé de l'offre"""
    top_skills = ", ".join(skills[:3])
    return (
        f"Offre de {title} {level} chez {company}. "
        f"Poste requérant {years} ans d'expérience minimum. "
        f"Compétences clés : {top_skills}. "
        f"Opportunité d'évolution dans un environnement technologique innovant."
    )


def generate_synthetic_job(index: int) -> dict:
    """Génère une fiche de poste synthétique complète"""

    # Choix du domaine et du titre
    domain_key = random.choice(list(JOBS.keys()))
    domain = JOBS[domain_key]
    title = random.choice(domain["titles"])

    # Informations générales
    company = random.choice(ENTREPRISES)
    industry = random.choice(domain["industries"])
    location = random.choice(VILLES)
    employment_type = random.choice(TYPES_EMPLOI)
    experience_level = random.choice(NIVEAUX_EXPERIENCE)
    education_req = random.choice(NIVEAUX_EDUCATION)
    years_required = get_experience_years(experience_level)

    # Compétences
    num_tech = random.randint(5, 10)
    num_langs = random.randint(2, 4)
    tech_skills = random.sample(domain["tech_skills"], min(num_tech, len(domain["tech_skills"])))
    prog_langs = random.sample(domain["prog_langs"], min(num_langs, len(domain["prog_langs"])))
    spoken_langs = random.sample(["Français", "Anglais", "Arabe"], random.randint(1, 3))

    # Certifications optionnelles
    certifications = []
    if random.random() > 0.6:
        certifications = random.sample(CERTIFICATIONS, random.randint(1, 2))

    # Séniorité
    top_techs = tech_skills[:3]
    top_langs = prog_langs[:2]
    seniority_tech = [{"technology": t, "level": experience_level} for t in top_techs]
    seniority_langs = [{"language": l, "level": experience_level} for l in top_langs]

    # Salaire
    salary_ranges = {
        "Junior": f"{random.randint(8, 12)}000-{random.randint(12, 15)}000 MAD/mois",
        "Medior": f"{random.randint(12, 18)}000-{random.randint(18, 25)}000 MAD/mois",
        "Confirmé": f"{random.randint(18, 25)}000-{random.randint(25, 35)}000 MAD/mois",
        "Senior": f"{random.randint(25, 35)}000-{random.randint(35, 50)}000 MAD/mois",
        "Expert": f"{random.randint(40, 60)}000-{random.randint(60, 80)}000 MAD/mois",
    }
    salary = salary_ranges.get(experience_level, "non specified")

    # Textes
    job_description = generate_job_description(
        title, tech_skills, prog_langs, experience_level, company, industry
    )
    summary = generate_summary(title, company, experience_level, tech_skills, years_required)

    # Date de publication
    months = ["2024-10", "2024-11", "2024-12", "2025-01", "2025-02", "2025-03"]
    days = [f"{random.randint(1, 28):02d}"]
    posted = f"{random.choice(months)}-{days[0]}"

    return {
        "index": index,
        "title": title,
        "company": company,
        "industry": industry,
        "location": location,
        "employment_type": employment_type,
        "job_description": job_description,
        "posted": posted,
        "programming_languages": prog_langs,
        "technical_skills": tech_skills,
        "spoken_languages": spoken_langs,
        "certifications": certifications,
        "seniority_requirements_technologies": seniority_tech,
        "seniority_requirements_programming_languages": seniority_langs,
        "experience_level": experience_level,
        "salary_range": salary,
        "education_requirements": education_req,
        "years_of_experience_required": years_required,
        "summary": summary
    }



def save_as_docx(job: dict, output_dir: str) -> str:
    """Génère un fichier DOCX réaliste pour la fiche de poste"""

    doc = Document()

    # Style du titre
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(f"{job['title']} — {job['experience_level']}")
    run.bold = True
    run.font.size = Pt(16)

    # Entreprise et localisation
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run(
        f"{job['company']} | {job['location']} | {job['employment_type']} | {job['posted']}"
    )
    info_run.font.size = Pt(11)

    doc.add_paragraph()

    # Section : Description du poste
    doc.add_heading("Description du poste", level=2)
    doc.add_paragraph(job["job_description"])

    # Section : Secteur et industrie
    doc.add_heading("Secteur d'activité", level=2)
    doc.add_paragraph(job["industry"])

    # Section : Compétences techniques
    doc.add_heading("Compétences techniques requises", level=2)
    tech_para = doc.add_paragraph()
    tech_para.add_run("Technologies : ").bold = True
    tech_para.add_run(", ".join(job["technical_skills"]))

    lang_para = doc.add_paragraph()
    lang_para.add_run("Langages de programmation : ").bold = True
    lang_para.add_run(", ".join(job["programming_languages"]))

    # Section : Langues
    doc.add_heading("Langues requises", level=2)
    doc.add_paragraph(", ".join(job["spoken_languages"]))

    # Section : Certifications
    if job["certifications"]:
        doc.add_heading("Certifications souhaitées", level=2)
        doc.add_paragraph(", ".join(job["certifications"]))

    # Section : Séniorité
    doc.add_heading("Niveau d'expérience requis", level=2)
    exp_para = doc.add_paragraph()
    exp_para.add_run("Niveau : ").bold = True
    exp_para.add_run(job["experience_level"])
    years_para = doc.add_paragraph()
    years_para.add_run("Années d'expérience minimum : ").bold = True
    years_para.add_run(str(job["years_of_experience_required"]))

    # Section : Formation
    doc.add_heading("Formation requise", level=2)
    doc.add_paragraph(job["education_requirements"])

    # Section : Salaire
    doc.add_heading("Package salarial", level=2)
    doc.add_paragraph(job["salary_range"])

    # Section : Résumé
    doc.add_heading("Résumé de l'offre", level=2)
    doc.add_paragraph(job["summary"])

    # Séniorité technique
    doc.add_heading("Profil de séniorité recherché", level=2)
    for item in job["seniority_requirements_technologies"]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{item['technology']} : ").bold = True
        p.add_run(item["level"])
    for item in job["seniority_requirements_programming_languages"]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{item['language']} : ").bold = True
        p.add_run(item["level"])

    # Sauvegarde
    title_clean = job["title"].replace(" ", "_").replace("/", "_").lower()
    company_clean = job["company"].replace(" ", "_").replace("&", "and").lower()
    filename = f"job_{job['index']:04d}_{title_clean}_{company_clean}.docx"
    filepath = os.path.join(output_dir, filename)
    os.makedirs(output_dir, exist_ok=True)
    doc.save(filepath)

    return filepath



def generate_dataset(num_jobs: int = 50, output_dir: str = "synthetic_jobs"):
    """
    Génère un dataset complet de fiches de poste synthétiques

    Args:
        num_jobs   : nombre de jobs à générer
        output_dir : dossier de sortie
    """
    print(f"\nGénération de {num_jobs} fiches de poste synthétiques...")
    print(f"Dossier de sortie : {output_dir}/")
    print("-" * 50)

    os.makedirs(os.path.join(output_dir, "docx"), exist_ok=True)

    all_jobs = []
    stats = {"total": num_jobs, "domains": {}, "levels": {}}

    for i in range(num_jobs):
        job = generate_synthetic_job(i + 1)
        all_jobs.append(job)

        # DOCX
        docx_path = save_as_docx(job, os.path.join(output_dir, "docx"))

        # Statistiques
        level = job["experience_level"]
        stats["levels"][level] = stats["levels"].get(level, 0) + 1

        if (i + 1) % 10 == 0:
            print(f"  {i + 1}/{num_jobs} jobs générés...")

    # Sauvegarde JSON de référence
    json_path = os.path.join(output_dir, "jobs_dataset.json")
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(all_jobs, f, ensure_ascii=False, indent=2)

    # Rapport
    print("\n" + "=" * 50)
    print("GÉNÉRATION TERMINÉE")
    print("=" * 50)
    print(f"  Total         : {num_jobs}")
    print(f"  Dossier DOCX  : {output_dir}/docx/")
    print(f"  Dataset JSON  : {json_path}")
    print("\n  Distribution par niveau :")
    for level in ["Junior", "Medior", "Confirmé", "Senior", "Expert"]:
        count = stats["levels"].get(level, 0)
        bar = "█" * count
        print(f"    {level:12} : {count:3}  {bar}")

    return all_jobs


if __name__ == "__main__":
    # Génère les jobs dans le bon dossier selon l'architecture CV_JOB/
    base_dir = Path(__file__).parent
    output = base_dir / "data" / "synthetic_data" / "jobs"

    jobs = generate_dataset(
        num_jobs=50,          # Change selon tes besoins
        output_dir=str(output)
    )

    print(f"\nExemple de job généré :")
    ex = jobs[0]
    print(f"  Titre      : {ex['title']}")
    print(f"  Entreprise : {ex['company']}")
    print(f"  Niveau     : {ex['experience_level']}")
    print(f"  Skills     : {', '.join(ex['technical_skills'][:3])}...")